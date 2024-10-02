from django.shortcuts import render

# Create your views here.
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework.parsers import MultiPartParser, FormParser
import spacy
import re

from docx import Document
import PyPDF2


class ResumeUploadAPI(APIView):
    parser_classes = (MultiPartParser, FormParser)

    def post(self, request, *args, **kwargs):
        resume_file = request.FILES['file']
        text_content = self.extract_text(resume_file)
        if resume_file:
            summary = self.summarize_resume(text_content)
            return Response({'summary': summary})
        else:
            return Response({"error": "No file uploaded"}, status=400)
       

    def extract_text(self, resume_file):
        if resume_file.name.endswith('.docx'):
            doc = Document(resume_file)
            return '\n'.join([para.text for para in doc.paragraphs])
        elif resume_file.name.endswith('.pdf'):
            pdf_reader = PyPDF2.PdfReader(resume_file)
            return '\n'.join([page.extract_text() for page in pdf_reader.pages])
        else:
            return ""

    # def summarize_resume(self, text_content):
    #     nlp = spacy.load('en_core_web_sm')
    #     doc = nlp(text_content)
    #     # NLP or keyword-based summarization logic here
    #     return "Summarized text"

    def summarize_resume(self, text_content):
         # Load the spaCy model for English language processing
        nlp = spacy.load('en_core_web_sm')
    
        # Process the text content using the NLP pipeline
        doc = nlp(text_content)
    
        # Initialize the summary dictionary
        summary = {
                "Name": None,
                "Email": None,
                "Phone": None,
                "Education": [],
                "Experience": [],
                "Skills": []
            }

        # Extract name from the first few lines
        lines = text_content.split('\n')
        for line in lines[:5]:  # Check only the first few lines for a name
                if len(line.split()) > 1 and not line.startswith('http'):
                    summary['Name'] = line
                    break

        # Use regex to extract email and phone number
        email_regex = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
        phone_regex = r'\+?\d[\d -]{8,12}\d'
    
        email_match = re.search(email_regex, text_content)
        phone_match = re.search(phone_regex, text_content)
    
        summary["Email"] = email_match.group() if email_match else None
        summary["Phone"] = phone_match.group() if phone_match else None
    
        # Extract entities and classify them
        for ent in doc.ents:
            if ent.label_ == "ORG" and any(keyword in ent.text.lower() for keyword in ['college', 'university']):
                summary["Education"].append(ent.text)
            elif ent.label_ == "ORG" and not any(keyword in ent.text.lower() for keyword in ['college',     'university']):
                summary["Experience"].append(ent.text)

        # A keyword-based approach to skill extraction
        skill_keywords = ['java', 'python', 'html', 'css', 'javascript', 'sql', 'springboot']
    
        # Token-based analysis for skill extraction
        tokens = [token.text.lower() for token in doc if not token.is_stop and not token.is_punct]
        for skill in skill_keywords:
            if skill in tokens:
                summary["Skills"].append(skill.capitalize())
    
        # Remove duplicates and clean up
        summary["Education"] = list(set(summary["Education"]))
        summary["Experience"] = list(set(summary["Experience"]))
        summary["Skills"] = list(set(summary["Skills"]))
    
        # Return the refined summary
        return summary


#--------------------------------------------------------------------------

# from django.shortcuts import render
# from rest_framework.views import APIView
# from rest_framework.response import Response
# from rest_framework.parsers import MultiPartParser, FormParser
# import spacy
# from docx import Document
# import PyPDF2

# class ResumeUploadAPI(APIView):
#     parser_classes = (MultiPartParser, FormParser)

#     # Load the spaCy model once when the class is initialized
#     def __init__(self, *args, **kwargs):
#         super().__init__(*args, **kwargs)
#         self.nlp = spacy.load('en_core_web_sm')

#     def post(self, request, *args, **kwargs):
#         resume_file = request.FILES.get('file')  # Safely retrieve the file

#         # Check if the file was provided in the request
#         if not resume_file:
#             return Response({"error": "No file uploaded"}, status=400)

#         # Extract text from the uploaded resume
#         try:
#             text_content = self.extract_text(resume_file)
#             if not text_content:
#                 return Response({"error": "Unsupported file format"}, status=400)
#         except Exception as e:
#             return Response({"error": f"Error processing file: {str(e)}"}, status=400)

#         # Summarize the extracted text
#         summary = self.summarize_resume(text_content)

#         return Response({'summary': summary, "message": "File uploaded and summarized successfully!"})

#     # Function to extract text from either .docx or .pdf files
#     def extract_text(self, resume_file):
#         # Handle .docx files
#         if resume_file.name.endswith('.docx'):
#             try:
#                 doc = Document(resume_file)
#                 return '\n'.join([para.text for para in doc.paragraphs])
#             except Exception as e:
#                 raise ValueError(f"Error processing .docx file: {str(e)}")
        
#         # Handle .pdf files
#         elif resume_file.name.endswith('.pdf'):
#             try:
#                 pdf_reader = PyPDF2.PdfReader(resume_file)
#                 return '\n'.join([page.extract_text() for page in pdf_reader.pages])
#             except Exception as e:
#                 raise ValueError(f"Error processing .pdf file: {str(e)}")
        
#         # If the file is neither .docx nor .pdf
#         else:
#             raise ValueError("Unsupported file format")

#     # Function to summarize the extracted text using spaCy
#     def summarize_resume(self, text_content):
#         doc = self.nlp(text_content)
#         # Implement your NLP-based summarization or keyword extraction logic here
#         # For now, this is just a placeholder
#         return "Summarized text (This is a placeholder, implement logic)"
